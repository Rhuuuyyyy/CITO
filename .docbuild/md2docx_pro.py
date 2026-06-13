#!/usr/bin/env python3
"""Conversor Markdown -> DOCX de nivel profissional (CITO).

Regra de paginacao de blocos cercados:
- Diagramas em ASCII (alta proporcao de caracteres de desenho) permanecem
  SEMPRE coesos (nunca partidos entre paginas).
- Trechos de codigo curtos (<= CODE_KEEP_MAX_LINES) permanecem coesos.
- Listagens de codigo longas podem fluir entre paginas (evita lacunas grandes).
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Emu

INK = RGBColor(0x0A, 0x0A, 0x0A)
MUTED = RGBColor(0x5A, 0x5A, 0x5A)
SOFT = RGBColor(0x8A, 0x8A, 0x8A)
ACCENT = RGBColor(0x1F, 0x1F, 0x1F)
CODE_INK = RGBColor(0x22, 0x22, 0x22)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CODE_BG = "F4F4F4"
HEADER_BG = "1F1F1F"
ZEBRA_BG = "F7F7F7"
BORDER = "D9D9D9"

BODY_FONT = "Calibri"
MONO_FONT = "Consolas"
BODY_SIZE = 10.5
CODE_KEEP_MAX_LINES = 8
# Caracteres tipicos de desenho/diagrama em ASCII.
ART_CHARS = set("+|│─-/\\<>v^▲▼►◄→←↑↓┌┐└┘├┤┬┴┼═║╔╗╚╝")
ART_RATIO_DIAGRAM = 0.22


def _is_diagram(lines):
    text = "".join(lines)
    nonspace = [c for c in text if not c.isspace()]
    if not nonspace:
        return False
    art = sum(1 for c in nonspace if c in ART_CHARS)
    return art / len(nonspace) > ART_RATIO_DIAGRAM


def _shade(el_props, color_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), color_hex)
    el_props.append(shd)


def _set_cell_bg(cell, color_hex):
    _shade(cell._tc.get_or_add_tcPr(), color_hex)


def _para_borders(p, edges, sz="4", color=BORDER, space="6"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    for edge in edges:
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), sz)
        e.set(qn("w:space"), space); e.set(qn("w:color"), color)
        pbdr.append(e)
    pPr.append(pbdr)


def _widow(p):
    pPr = p._p.get_or_add_pPr()
    w = OxmlElement("w:widowControl"); w.set(qn("w:val"), "true"); pPr.append(w)


def _add_runs(paragraph, text, size=BODY_SIZE, color=INK, bold=False, italic=False):
    pattern = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(size); r.font.color.rgb = color
            r.font.bold = bold; r.font.italic = italic; r.font.name = BODY_FONT
        tok = m.group(0)
        if tok.startswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.font.bold = True
            r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = BODY_FONT
        elif tok.startswith("`"):
            r = paragraph.add_run(tok[1:-1])
            r.font.name = MONO_FONT; r.font.size = Pt(size - 1); r.font.color.rgb = ACCENT
            _shade(r._element.get_or_add_rPr(), CODE_BG)
        else:
            r = paragraph.add_run(tok[1:-1]); r.font.italic = True
            r.font.size = Pt(size); r.font.color.rgb = color; r.font.name = BODY_FONT
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.bold = bold; r.font.italic = italic; r.font.name = BODY_FONT


def _heading(doc, level, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.keep_with_next = True; pf.keep_together = True
    pf.space_before = Pt(18 if level == 1 else 12 if level == 2 else 9)
    pf.space_after = Pt(7 if level <= 2 else 4)
    _widow(p)
    sizes = {1: 17, 2: 13.5, 3: 11.5, 4: 11}
    color = INK if level <= 2 else ACCENT
    _add_runs(p, text, size=sizes.get(level, 11), color=color, bold=True)
    if level <= 2:
        _para_borders(p, ["bottom"], sz=("8" if level == 1 else "4"),
                      color=("1F1F1F" if level == 1 else "BFBFBF"), space="4")


def _paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(7)
    _widow(p)
    _add_runs(p, text)


def _code_block(doc, lines):
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.12); pf.right_indent = Inches(0.05)
    pf.space_before = Pt(4); pf.space_after = Pt(9)
    # Diagramas nunca partem; codigo curto fica coeso; codigo longo flui.
    if _is_diagram(lines):
        pf.keep_together = True
    else:
        pf.keep_together = (len(lines) <= CODE_KEEP_MAX_LINES)
    pf.line_spacing = 1.0
    _shade(p._p.get_or_add_pPr(), CODE_BG)
    _para_borders(p, ["top", "bottom", "left", "right"], sz="4", color="DDDDDD", space="6")
    run = p.add_run("\n".join(lines))
    run.font.name = MONO_FONT; run.font.size = Pt(8.5); run.font.color.rgb = CODE_INK


def _quote(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.28); pf.space_before = Pt(4); pf.space_after = Pt(9)
    pf.keep_together = True
    _para_borders(p, ["left"], sz="18", color="1F1F1F", space="10")
    _add_runs(p, text, size=BODY_SIZE, color=MUTED, italic=True)


def _list_item(doc, text, ordered, idx, level=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Inches(0.32 + 0.28 * level)
    pf.first_line_indent = Inches(-0.20)
    pf.space_after = Pt(3)
    _widow(p)
    r = p.add_run((f"{idx}. " if ordered else "•\t"))
    r.font.size = Pt(BODY_SIZE); r.font.bold = ordered
    r.font.color.rgb = ACCENT; r.font.name = BODY_FONT
    _add_runs(p, text)


def _col_widths(rows, total_emu):
    ncols = max(len(r) for r in rows)
    maxlen = [1] * ncols
    for r in rows:
        for j in range(ncols):
            cell = r[j] if j < len(r) else ""
            longest = max((len(w) for w in cell.split()), default=1)
            maxlen[j] = max(maxlen[j], int(len(cell) * 0.5 + longest))
    floor = 6
    weights = [max(floor, m) for m in maxlen]
    s = sum(weights)
    widths = [int(total_emu * w / s) for w in weights]
    widths[-1] += total_emu - sum(widths)
    return widths


def _add_table(doc, rows, content_emu):
    header, *body = rows
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False; table.allow_autofit = False

    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout"); layout.set(qn("w:type"), "fixed"); tblPr.append(layout)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa"); tblW.set(qn("w:w"), str(int(content_emu / 635))); tblPr.append(tblW)
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", "40"), ("bottom", "40"), ("left", "90"), ("right", "90")):
        e = OxmlElement(f"w:{edge}"); e.set(qn("w:w"), val); e.set(qn("w:type"), "dxa"); mar.append(e)
    tblPr.append(mar)

    widths = _col_widths(rows, content_emu)
    widths_dxa = [max(1, int(w / 635)) for w in widths]
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), widths_dxa):
            gc.set(qn("w:w"), str(w))

    def _apply_widths(cells):
        for j, c in enumerate(cells):
            c.width = Emu(widths_dxa[j] * 635)

    hdr = table.rows[0]
    trPr = hdr._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:tblHeader"))
    trPr.append(OxmlElement("w:cantSplit"))
    for j, txt in enumerate(header):
        c = hdr.cells[j]
        _set_cell_bg(c, HEADER_BG)
        para = c.paragraphs[0]
        para.paragraph_format.space_after = Pt(1); para.paragraph_format.space_before = Pt(1)
        _add_runs(para, txt, size=9.5, color=WHITE, bold=True)
    _apply_widths(hdr.cells)

    for i, row in enumerate(body):
        r = table.add_row()
        r._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        cells = r.cells
        if i % 2 == 1:
            for c in cells:
                _set_cell_bg(c, ZEBRA_BG)
        for j in range(ncols):
            val = row[j] if j < len(row) else ""
            para = cells[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(2); para.paragraph_format.space_before = Pt(1)
            _add_runs(para, val, size=9.5)
        _apply_widths(cells)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2); sp.paragraph_format.space_before = Pt(0)
    return table


def _parse_table(block):
    rows = []
    for ln in block:
        ln = ln.strip()
        if not ln.startswith("|"):
            continue
        rows.append([c.strip() for c in ln.strip().strip("|").split("|")])
    cleaned = []
    for r in rows:
        if all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in r if c != ""):
            continue
        cleaned.append(r)
    return cleaned


def _cover(doc, title, subtitle, paras):
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title); r.font.bold = True; r.font.size = Pt(30)
    r.font.color.rgb = INK; r.font.name = BODY_FONT
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.paragraph_format.space_after = Pt(14)
    _para_borders(d, ["bottom"], sz="6", color="1F1F1F", space="1")
    if subtitle:
        s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(26)
        _add_runs(s, subtitle, size=14, color=MUTED)
    for i, txt in enumerate(paras):
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        last = (i == len(paras) - 1)
        pp.paragraph_format.space_after = Pt(8)
        pp.paragraph_format.left_indent = Inches(0.6); pp.paragraph_format.right_indent = Inches(0.6)
        if last and ("Versão" in txt or ("·" in txt and "Banco" in txt)):
            r = pp.add_run(txt); r.font.name = MONO_FONT; r.font.size = Pt(9); r.font.color.rgb = SOFT
        else:
            _add_runs(pp, txt, size=11, color=RGBColor(0x33, 0x33, 0x33))
    doc.add_page_break()


def _footer(section, doc_title):
    footer = section.footer
    p = footer.paragraphs[0]; p.text = ""
    content_w = section.page_width - section.left_margin - section.right_margin
    p.paragraph_format.tab_stops.add_tab_stop(content_w, WD_TAB_ALIGNMENT.RIGHT)
    _para_borders(p, ["top"], sz="4", color="D9D9D9", space="4")
    rl = p.add_run(doc_title); rl.font.size = Pt(8); rl.font.color.rgb = SOFT; rl.font.name = BODY_FONT
    rt = p.add_run("\t"); rt.font.size = Pt(8); rt.font.color.rgb = SOFT
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    rp = p.add_run(); rp.font.size = Pt(8); rp.font.color.rgb = MUTED
    rp._r.append(fld1); rp._r.append(instr); rp._r.append(fld2)


def convert(md_path, docx_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0); section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0); section.right_margin = Inches(1.0)
    content_emu = int(section.page_width) - int(section.left_margin) - int(section.right_margin)

    style = doc.styles["Normal"]
    style.font.name = BODY_FONT; style.font.size = Pt(BODY_SIZE); style.font.color.rgb = INK
    pf = style.paragraph_format
    pf.space_after = Pt(7); pf.line_spacing = 1.18; pf.widow_control = True

    title = ""; subtitle = ""; cover_paras = []
    i = 0; n = len(lines)
    while i < n and not lines[i].strip():
        i += 1
    m = re.match(r"^#\s+(.*)$", lines[i].strip()) if i < n else None
    if m:
        title = m.group(1).strip(); i += 1
    while i < n:
        s = lines[i].strip()
        if s == "<<<PAGEBREAK>>>":
            i += 1; break
        if re.match(r"^#\s+\d", s):
            break
        if re.fullmatch(r"-{3,}", s):
            break
        m2 = re.match(r"^##\s+(.*)$", s)
        if m2 and not subtitle and not cover_paras:
            subtitle = m2.group(1).strip(); i += 1; continue
        if re.match(r"^#{1,6}\s", s):
            break
        if s:
            cover_paras.append(s)
        i += 1

    _footer(section, title)
    _cover(doc, title, subtitle, cover_paras)

    while i < n:
        line = lines[i]; s = line.strip()
        if not s:
            i += 1; continue
        if s == "<<<PAGEBREAK>>>":
            doc.add_page_break(); i += 1; continue
        if s.startswith("```"):
            j = i + 1; buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j]); j += 1
            _code_block(doc, buf); i = j + 1; continue
        mh = re.match(r"^(#{1,4})\s+(.*)$", s)
        if mh:
            _heading(doc, len(mh.group(1)), mh.group(2).strip()); i += 1; continue
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", s):
            i += 1; continue
        if s.startswith("|"):
            j = i; block = []
            while j < n and lines[j].strip().startswith("|"):
                block.append(lines[j]); j += 1
            rows = _parse_table(block)
            if rows:
                _add_table(doc, rows, content_emu)
            i = j; continue
        if s.startswith(">"):
            qbuf = []
            while i < n and lines[i].strip().startswith(">"):
                qbuf.append(lines[i].strip()[1:].strip()); i += 1
            _quote(doc, " ".join(qbuf)); continue
        mo = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if mo:
            _list_item(doc, mo.group(3).strip(), True, mo.group(2), level=len(mo.group(1)) // 2)
            i += 1; continue
        mb = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if mb:
            _list_item(doc, mb.group(2).strip(), False, None, level=len(mb.group(1)) // 2)
            i += 1; continue
        para = [s]; i += 1
        while i < n and lines[i].strip() and not re.match(
            r"^(#{1,4}\s|\||>|```|\s*[-*]\s|\s*\d+\.\s)", lines[i]
        ) and lines[i].strip() not in ("---", "<<<PAGEBREAK>>>"):
            para.append(lines[i].strip()); i += 1
        _paragraph(doc, " ".join(para))

    doc.save(docx_path)
    print(f"OK: {docx_path}")


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
